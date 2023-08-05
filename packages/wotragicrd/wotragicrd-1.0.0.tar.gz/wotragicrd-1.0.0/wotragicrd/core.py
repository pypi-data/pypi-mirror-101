from docx import Document


def file():
    document: object = Document()
    document.add_heading('Word file from package', 0)

    document.add_paragraph('Its table').italic = True

    table = document.add_table(rows=3, cols=3)
    c = table.rows[0].cells
    c[0].text = 't'
    c[1].text = 'h'
    c[2].text = 'i'
    f = table.rows[1].cells
    f[0].text = 's'
    f[1].text = 'i'
    f[2].text = 's'
    g = table.rows[2].cells
    g[0].text = 't'
    g[1].text = 'a'
    g[2].text = 'b'

    document.save('jene4ka.docx')